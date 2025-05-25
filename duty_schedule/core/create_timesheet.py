from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm
import os
from datetime import datetime, timedelta


def create_timesheet(user_data, attendance_data, start_date, end_date):
    # Создаем новый документ
    doc = Document()

    # Устанавливаем альбомную ориентацию
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # Устанавливаем шрифт по умолчанию
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(9)

    # Первая таблица (коды, сужена по тексту и прижата к правому краю)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.autofit = True  # Сужаем по содержимому
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cells = table.rows[0].cells
    cells[0].text = 'Код'
    cells[1].text = 'Форма по ОКУД'
    cells[2].text = '0301008'

    cells = table.rows[1].cells
    cells[1].text = 'по ОКПО'

    cells = table.rows[2].cells
    cells[0].text = 'УМВД России по Брянской области'
    cells[0].merge(table.rows[2].cells[2])

    cells = table.rows[3].cells
    cells[0].text = 'Вычислительный центр ИЦ'
    cells[0].merge(table.rows[3].cells[2])

    # Вторая таблица (номер документа и период, сужена по тексту и прижата к правому краю)
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    table.autofit = True  # Сужаем по содержимому
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cells = table.rows[0].cells
    cells[0].text = 'Номер документа'
    cells[1].text = 'Дата составления'
    cells[2].text = 'Отчетный период'
    cells[2].merge(table.rows[0].cells[4])  # Объединяем до конца строки

    cells = table.rows[1].cells
    cells[3].text = start_date.strftime('%d.%m.%Y')
    cells[4].text = end_date.strftime('%d.%m.%Y')

    # Заголовок табеля (по центру)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ТАБЕЛЬ\nучета рабочего времени')
    run.bold = True

    # Основная таблица с данными
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    # Заголовки
    headers = ['ФИО', 'Дата', 'Отметка', 'Часы']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Заполняем данные
    current_row = 1
    current_date = start_date
    total_hours = 0  # Счетчик общего количества часов

    while current_date <= end_date:
        # Добавляем новую строку при необходимости
        if current_row >= len(table.rows):
            table.add_row()
        
        row = table.rows[current_row]
        # ФИО
        row.cells[0].text = f"{user_data['last_name']} {user_data['first_name']} {user_data['middle_name']}".strip()
        # Дата
        row.cells[1].text = current_date.strftime('%d.%m.%Y')
        # Отметка
        is_present = attendance_data.get(current_date.date(), False)
        row.cells[2].text = 'Я' if is_present else 'Н'
        # Часы за день
        hours = 8 if is_present else 0
        row.cells[3].text = str(hours)
        total_hours += hours
        
        current_row += 1
        current_date += timedelta(days=1)

    # Добавляем итоговую строку
    summary_row = table.add_row()
    summary_row.cells[0].text = "Итого:"
    summary_row.cells[0].merge(summary_row.cells[2])  # Объединяем ячейки от ФИО до Отметки
    summary_row.cells[3].text = str(total_hours)  # Общее количество часов
    
    # Выделяем итоговую строку жирным шрифтом
    for cell in summary_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Подписи
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('Работник кадровой службы          ______________________________________           ______________________________________')
    p.add_run('\n                                                                                личная подпись                                                       расшифровка подписи')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p.add_run('\n«____» ________________ 20__ года')

    return doc


if __name__ == '__main__':
    user_data = {
        'last_name': 'Иванов',
        'first_name': 'Иван',
        'middle_name': 'Иванович'
    }
    attendance_data = {
        datetime(2024, 1, 1).date(): True,
        datetime(2024, 1, 2).date(): True,
        datetime(2024, 1, 3).date(): False
    }
    start_date = datetime(2024, 1, 1)
    end_date = datetime(2024, 1, 3)
    doc = create_timesheet(user_data, attendance_data, start_date, end_date)
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    file_path = os.path.join(desktop_path, "T13_new.docx")
    doc.save(file_path)
    print(f"Файл сохранен по пути: {file_path}") 